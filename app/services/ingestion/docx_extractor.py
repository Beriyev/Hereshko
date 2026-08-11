from docx import Document as DocxDocument
from app.core.normalization import Document, SourceType
import uuid
from pathlib import Path
from datetime import datetime, timezone
from app.core.exceptions import IngestionError

def extract_docx(file_path: Path, notebook_id: str) -> Document:
    texts = []
    boundaries = []
    offset = 0
    para_number = 1
    paras_count = 0
    try:
        doc = DocxDocument(str(file_path))
        paras_count = len(doc.paragraphs)  
        for paragraph in doc.paragraphs:
            text = paragraph.text
            text_size = len(text) if text else 0
            if text:
                boundaries.append({
                    "paragraph_index": para_number,
                    "start": offset,
                    "end": offset + text_size
                })
                texts.append(text)
                offset+= text_size+1
            para_number += 1
    except Exception as e:
        raise IngestionError(f"Failed to read DOCX file: {e}") from e
    content = "\n".join(texts)+"\n" if texts else ""
    if not content.strip():
        raise IngestionError("No extractable text found in DOCX (may be an empty document)")
    return Document(
        document_id = str(uuid.uuid4()),
        notebook_id = notebook_id,
        content = content,
        source_type = SourceType.DOCX,
        source_identifier = file_path.name,
        title = file_path.stem,
        ingested_at = datetime.now(timezone.utc),
        raw_metadata = {
            "original_filename": file_path.name,
            "file_size_bytes": file_path.stat().st_size,
            "paras_count": paras_count,
            "boundaries": boundaries
        }
    )
    


